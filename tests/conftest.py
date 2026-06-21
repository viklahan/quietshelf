"""Test fixtures. Environment is set before the app is imported."""
from __future__ import annotations

import os

import pytest

# Must be set before app modules are imported by any test.
os.environ["LLM_PROVIDER"] = "gemini"
os.environ["GEMINI_API_KEY"] = "test-key-not-real"
os.environ.setdefault("RATE_LIMIT", "1000")  # don't trip the limiter in tests
os.environ.pop("ACCESS_CODE", None)
os.environ.pop("MODEL_NAME", None)

VALID_SCRIPT = ("the quick brown fox jumps over the lazy dog again and again " * 15).strip()

VALID_SHOT_LIST = {
    "video_title_suggestion": "A Test Title",
    "estimated_runtime_seconds": 60,
    "segments": [
        {
            "id": 1,
            "script_text": "the quick brown fox jumps over the lazy dog.",
            "start_time": "0:00",
            "end_time": "0:09",
            "search_terms": ["fox running field", "dog sleeping grass", "forest morning light"],
            "clip_duration_seconds": 9,
            "mood": "playful",
        }
    ],
}


@pytest.fixture()
def client():
    from fastapi.testclient import TestClient

    from app.main import app

    with TestClient(app) as test_client:
        yield test_client


@pytest.fixture()
def valid_script() -> str:
    return VALID_SCRIPT


@pytest.fixture()
def valid_shot_list() -> dict:
    return VALID_SHOT_LIST


@pytest.fixture()
def sample_docx(tmp_path):
    """A small real DOCX with a heading and paragraphs, for Format tests."""
    from docx import Document

    doc = Document()
    doc.add_heading("Chapter One", level=1)
    doc.add_paragraph("It was a bright cold day and the clocks were striking.")
    doc.add_paragraph("She walked the long road home, thinking of nothing at all.")
    doc.add_heading("Chapter Two", level=1)
    doc.add_paragraph("The second chapter opens on a quiet morning by the sea.")
    path = tmp_path / "manuscript.docx"
    doc.save(path)
    return path
