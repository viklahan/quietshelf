#!/usr/bin/env python3
"""
Quiet Shelf end-to-end smoke test.
Run from Legion7: python tests/qs_e2e_test.py  (add --ai to include AI tests)
Tests every tab's browse/paste path against the live site.
"""
import sys
import io
import time
import json
import random
import zipfile
import tempfile
import os
from pathlib import Path

try:
    import requests
except ImportError:
    print("INSTALL: pip install requests")
    sys.exit(1)

BASE = "https://quietshelf.studio"
# AI-consuming tests (blurb synthesis, promote mapping) are OPT-IN: they spend
# real provider quota on the server, which is production capacity. Default run
# costs zero AI calls. Enable with:  python tests/qs_e2e_test.py --ai
AI_TESTS = "--ai" in sys.argv
PASS = 0
FAIL = 0
RESULTS = []

def ok(name, detail=""):
    global PASS
    PASS += 1
    RESULTS.append(("PASS", name, detail))
    print(f"  ✓  {name}" + (f" — {detail}" if detail else ""))

def fail(name, detail=""):
    global FAIL
    FAIL += 1
    RESULTS.append(("FAIL", name, detail))
    print(f"  ✗  {name}" + (f" — {detail}" if detail else ""))

def section(title):
    print(f"\n{'='*50}")
    print(f"  {title}")
    print('='*50)

# ── Test content ────────────────────────────────────────────────────────────

SENTENCES = [
    "The old lighthouse stood at the edge of the world, its beam cutting through the fog like a knife through silk.",
    "Margaret had lived on the island for thirty years, and she still felt like a stranger to its moods.",
    "The fishermen came back with empty nets again, their faces drawn with the particular exhaustion of men who have failed the sea.",
    "She found the letter tucked beneath the door, written in a hand she had not seen since childhood.",
    "The storm had been building for three days before it finally broke over the harbor like a wave of glass.",
    "Thomas knew the moment he saw the light in her window that everything was about to change.",
    "The captain had sailed these waters for forty years and had never seen anything like what lay ahead.",
    "Children ran through the narrow streets of the old town, their laughter echoing off cobblestones worn smooth by centuries.",
    "The doctor looked at his hands and wondered how many lives they had held and released.",
    "Night came slowly to the village, painting the sky in shades of amber and deep violet.",
    "She remembered her grandmother telling her that the sea remembers everyone who has ever drowned in it.",
    "The market was full of voices and color, a chaos that somehow resolved itself into order every morning.",
    "He had not spoken to his brother in seven years, and now there was nothing left to say.",
    "The old book fell open to a page she had never read, its words seeming to glow in the lamplight.",
    "Rain came in sideways off the Atlantic, turning the windows into rivers of moving glass.",
    "The child asked why the stars moved and her father said they were trying to find their way home.",
    "A fire burned low in the hearth, casting long shadows across the walls of the empty room.",
    "She packed slowly, folding each memory as carefully as a piece of cloth that might be needed later.",
    "The boat left the harbor at dawn, while the town still slept and the gulls had not yet begun their morning arguments.",
    "He opened the door and found only silence where a life had been.",
]

CHAPTERS = ["Prologue", "Chapter One: The Beginning", "Chapter Two: The Journey",
            "Chapter Three: The Storm", "Chapter Four: The Truth",
            "Chapter Five: The Reckoning", "Chapter Six: The Return", "Epilogue"]

def make_story_text(target_words=8000):
    lines = []
    for title in CHAPTERS:
        lines.append(title)
        lines.append("")
        for _ in range(15):
            para = " ".join(random.choices(SENTENCES, k=random.randint(5, 8)))
            lines.append(para)
            lines.append("")
        if len(" ".join(lines).split()) >= target_words:
            break
    return "\n".join(lines)

def make_docx(text, path):
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("The Lighthouse at the Edge of the World", 0)
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                continue
            if any(line.startswith(h.split(":")[0]) for h in CHAPTERS):
                doc.add_heading(line, level=1)
            else:
                doc.add_paragraph(line)
        doc.save(path)
        return True
    except ImportError:
        return False

STORY_8K = make_story_text(8000)
STORY_5K = " ".join(STORY_8K.split()[:5000])
STORY_1K = " ".join(STORY_8K.split()[:1200])
STORY_100 = " ".join(STORY_8K.split()[:150])

print(f"\nTest content: 8K={len(STORY_8K.split())}w  5K={len(STORY_5K.split())}w  100={len(STORY_100.split())}w")

# ── 1. Site health ───────────────────────────────────────────────────────────
section("1. SITE HEALTH")

try:
    r = requests.get(BASE + "/", timeout=10)
    if r.status_code == 200:
        ok("Site loads", f"HTTP {r.status_code}")
    else:
        fail("Site loads", f"HTTP {r.status_code}")
except Exception as e:
    fail("Site loads", str(e))

try:
    r = requests.get(BASE + "/api/format/themes", timeout=10)
    themes = r.json().get("themes", [])
    if len(themes) == 4:
        ok("Format themes", f"{len(themes)} themes")
    else:
        fail("Format themes", f"got {len(themes)}")
except Exception as e:
    fail("Format themes", str(e))

try:
    r = requests.get(BASE + "/api/health", timeout=10)
    d = r.json()
    provider = d.get("provider", "unknown")
    ok("Health endpoint", f"provider={provider}")
except Exception as e:
    fail("Health endpoint", str(e))

# ── 2. FORMAT TAB ────────────────────────────────────────────────────────────
section("2. FORMAT TAB")

with tempfile.TemporaryDirectory() as tmp:
    # Test TXT upload
    txt_path = os.path.join(tmp, "story.txt")
    with open(txt_path, "w") as f:
        f.write(STORY_8K)
    try:
        t0 = time.time()
        r = requests.post(BASE + "/api/format",
            files={"file": ("story.txt", open(txt_path, "rb"), "text/plain")},
            data={"title": "The Lighthouse", "author": "Test Author", "theme": "classic"},
            timeout=60)
        elapsed = round(time.time() - t0, 1)
        if r.status_code == 200 and zipfile.is_zipfile(io.BytesIO(r.content)):
            ok("Format TXT (8K words)", f"{len(r.content)//1024}KB in {elapsed}s")
        else:
            body = r.text[:200] if r.status_code != 200 else "invalid EPUB"
            fail("Format TXT (8K words)", f"HTTP {r.status_code}: {body}")
    except Exception as e:
        fail("Format TXT (8K words)", str(e))

    # Test DOCX upload
    docx_path = os.path.join(tmp, "story.docx")
    if make_docx(STORY_8K, docx_path):
        try:
            t0 = time.time()
            r = requests.post(BASE + "/api/format",
                files={"file": ("story.docx", open(docx_path, "rb"),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"title": "The Lighthouse", "author": "Test Author", "theme": "cozy"},
                timeout=60)
            elapsed = round(time.time() - t0, 1)
            if r.status_code == 200 and zipfile.is_zipfile(io.BytesIO(r.content)):
                ok("Format DOCX (8K words)", f"{len(r.content)//1024}KB in {elapsed}s")
            else:
                fail("Format DOCX (8K words)", f"HTTP {r.status_code}: {r.text[:200]}")
        except Exception as e:
            fail("Format DOCX (8K words)", str(e))
    else:
        print("  SKIP  Format DOCX (python-docx not installed)")

    # Test empty author fallback
    try:
        r = requests.post(BASE + "/api/format",
            files={"file": ("story.txt", open(txt_path, "rb"), "text/plain")},
            data={"title": "Test", "author": "", "theme": "modern"},
            timeout=60)
        if r.status_code == 200:
            ok("Format empty author fallback", "no 422")
        else:
            fail("Format empty author fallback", f"HTTP {r.status_code}")
    except Exception as e:
        fail("Format empty author fallback", str(e))

# ── 3. BLURB TAB ─────────────────────────────────────────────────────────────
section("3. BLURB TAB")

if not AI_TESTS:
    print("  SKIP  Blurb tests (AI quota) — run with --ai to include")
else:
    # ONE call, one tone: the smoke question is "does blurb work", and every
    # tone rides the same code path. Four tones was quota gluttony.
    try:
        t0 = time.time()
        r = requests.post(BASE + "/api/blurb",
            data={"text": STORY_1K, "tone": "warm", "length": "medium"},
            timeout=60)
        elapsed = round(time.time() - t0, 1)
        if r.status_code == 200:
            d = r.json()
            has_blurb = bool(d.get("back_cover") or (d.get("back_cover_variants") and d["back_cover_variants"][0]))
            if has_blurb:
                ok("Blurb tone=warm", f"{elapsed}s")
            else:
                fail("Blurb tone=warm", "no back_cover in response")
        else:
            fail("Blurb tone=warm", f"HTTP {r.status_code}: {r.text[:150]}")
    except Exception as e:
        fail("Blurb tone=warm", str(e))

# ── 4. PROMOTE TAB ───────────────────────────────────────────────────────────
section("4. PROMOTE TAB")

# Test extract endpoint (file browse)
with tempfile.TemporaryDirectory() as tmp:
    txt_path = os.path.join(tmp, "story.txt")
    with open(txt_path, "w") as f:
        f.write(STORY_8K)
    try:
        r = requests.post(BASE + "/api/promote/extract",
            files={"file": ("story.txt", open(txt_path, "rb"), "text/plain")},
            timeout=30)
        if r.status_code == 200:
            d = r.json()
            wc = d.get("word_count", 0)
            ok("Promote extract TXT", f"{wc} words extracted")
        else:
            fail("Promote extract TXT", f"HTTP {r.status_code}: {r.text[:150]}")
    except Exception as e:
        fail("Promote extract TXT", str(e))

    docx_path = os.path.join(tmp, "story.docx")
    if make_docx(STORY_8K, docx_path):
        try:
            r = requests.post(BASE + "/api/promote/extract",
                files={"file": ("story.docx", open(docx_path, "rb"),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                timeout=30)
            if r.status_code == 200:
                d = r.json()
                wc = d.get("word_count", 0)
                ok("Promote extract DOCX", f"{wc} words extracted")
            else:
                fail("Promote extract DOCX", f"HTTP {r.status_code}: {r.text[:150]}")
        except Exception as e:
            fail("Promote extract DOCX", str(e))

# Test streaming (paste path) — AI-hungry: every chunk is a provider call
if not AI_TESTS:
    print("  SKIP  Promote stream (AI quota) — run with --ai to include")
else:
  try:
    t0 = time.time()
    r = requests.post(BASE + "/api/promote/stream",
        json={"script": STORY_1K},
        headers={"Accept": "text/event-stream"},
        stream=True,
        timeout=180)
    if r.status_code != 200:
        fail("Promote stream 1K", f"HTTP {r.status_code}: {r.text[:150]}")
    else:
        chunks_received = 0
        total_chunks = 0
        segments_total = 0
        got_done = False
        errors = []
        for line in r.iter_lines(decode_unicode=True):
            if not line or not line.startswith("data: "):
                continue
            try:
                evt = json.loads(line[6:])
                if evt["type"] == "meta":
                    total_chunks = evt.get("total_chunks", 0)
                elif evt["type"] == "chunk":
                    chunks_received += 1
                    segments_total += len(evt.get("segments", []))
                elif evt["type"] == "done":
                    got_done = True
                    break
                elif evt["type"] == "error":
                    errors.append(evt.get("message", "unknown"))
            except Exception:
                pass
        elapsed = round(time.time() - t0, 1)
        if errors:
            fail("Promote stream 1K", f"errors: {errors}")
        elif not got_done:
            fail("Promote stream 1K", f"no done event — got {chunks_received}/{total_chunks} chunks, {segments_total} segs")
        elif chunks_received < total_chunks:
            fail("Promote stream 1K", f"incomplete: {chunks_received}/{total_chunks} chunks, {segments_total} segs in {elapsed}s")
        else:
            ok("Promote stream 1K", f"{chunks_received}/{total_chunks} chunks, {segments_total} segs in {elapsed}s")
  except Exception as e:
    fail("Promote stream 1K", str(e))

# ── 5. COVER SUGGESTIONS ────────────────────────────────────────────────────
section("5. COVER SUGGESTIONS")

try:
    r = requests.post(BASE + "/api/format/cover-suggestions",
        data={"title": "The Lighthouse", "passage": STORY_100, "n": 3},
        timeout=30)
    if r.status_code == 200:
        suggestions = r.json().get("suggestions", [])
        if suggestions:
            sources = [s.get("source") for s in suggestions]
            ok("Cover suggestions", f"{len(suggestions)} photos from {sources}")
        else:
            ok("Cover suggestions", "0 photos (image API keys may not be set)")
    else:
        fail("Cover suggestions", f"HTTP {r.status_code}: {r.text[:150]}")
except Exception as e:
    fail("Cover suggestions", str(e))


# ── 6. SCOUT TAB ────────────────────────────────────────────────────────────
section("6. SCOUT TAB")

try:
    t0 = time.time()
    r = requests.post(BASE + "/api/scout/harvest",
        json={"sources": [], "seeds": ["feeling stuck"]},
        timeout=90)
    elapsed = round(time.time() - t0, 1)
    if r.status_code == 200:
        d = r.json()
        sc = d.get("suggestion_count", 0)
        wc = d.get("word_count", 0)
        if sc > 0:
            ok("Scout harvest (seeds only)", f"{sc} suggestions, {wc} words in {elapsed}s")
        else:
            fail("Scout harvest (seeds only)", "200 but zero suggestions")
    else:
        # NOTE: suggestion engines may block datacenter IPs. A 502 here with
        # engine errors is the server telling the truth about its network.
        fail("Scout harvest (seeds only)", f"HTTP {r.status_code}: {r.text[:200]}")
except Exception as e:
    fail("Scout harvest (seeds only)", str(e))

try:
    r = requests.post(BASE + "/api/scout/harvest", json={"sources": [], "seeds": []}, timeout=15)
    if r.status_code == 422:
        ok("Scout empty-input rejection", "422 as designed")
    else:
        fail("Scout empty-input rejection", f"expected 422, got {r.status_code}")
except Exception as e:
    fail("Scout empty-input rejection", str(e))

# ── SUMMARY ─────────────────────────────────────────────────────────────────
print(f"\n{'='*50}")
print(f"  RESULTS:  {PASS} passed  {FAIL} failed" + ("" if AI_TESTS else "  (AI tests skipped — use --ai)"))
print('='*50)
if FAIL > 0:
    print("\nFAILED:")
    for status, name, detail in RESULTS:
        if status == "FAIL":
            print(f"  ✗  {name}: {detail}")
sys.exit(0 if FAIL == 0 else 1)
